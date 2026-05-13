#!/usr/bin/env python3
"""check_resume_format.py — see what an ATS would actually extract from your PDF/DOCX.

Real ATSes parse PDF/DOCX with simple text extractors that frequently mangle
tables, multi-column layouts, and unusual fonts. This script extracts text
from your resume.pdf or resume.docx using a deliberately simple parser
(PyPDF2 / python-docx — no fancy layout reconstruction), compares it to your
clean resume.md, and warns about anything that would be lost or garbled.

Run this once when you update your actual PDF/DOCX resume. Doesn't affect
pipeline scoring — those scores always use resume.md.

Usage:
    python check_resume_format.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).parent
RESUME_MD = ROOT / "resume.md"
RESUME_PDF = ROOT / "resume.pdf"
RESUME_DOCX = ROOT / "resume.docx"


def _bar(label: str = "", width: int = 78) -> str:
    if not label:
        return "─" * width
    label = f" {label} "
    pad = max(0, width - len(label) - 4)
    return f"\n──{label}{'─' * pad}\n"


def extract_pdf_text(pdf_path: Path) -> str:
    """Extract text from a PDF using PyPDF2 (simple, ATS-like extraction)."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        sys.exit(
            "Error: PyPDF2 is not installed. Run `pip install PyPDF2` "
            "(or re-run pip install -r requirements.txt)."
        )
    reader = PdfReader(str(pdf_path))
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            print(f"  ! page {i+1}: extraction error — {exc}", file=sys.stderr)
            text = ""
        parts.append(text)
    return "\n".join(parts)


def extract_docx_text(docx_path: Path) -> str:
    """Extract text from a DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        sys.exit(
            "Error: python-docx is not installed. Run `pip install python-docx` "
            "(or re-run pip install -r requirements.txt)."
        )
    doc = Document(str(docx_path))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # Tables get extracted separately — and if they're present, that's a warning.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def find_resume_file() -> tuple[Path, str]:
    """Return (path, kind) for whichever resume file exists. PDF preferred over DOCX."""
    if RESUME_PDF.exists():
        return RESUME_PDF, "pdf"
    if RESUME_DOCX.exists():
        return RESUME_DOCX, "docx"
    sys.exit(
        f"Error: neither {RESUME_PDF.name} nor {RESUME_DOCX.name} found in {ROOT}. "
        f"Drop your real PDF or DOCX resume in this directory and re-run."
    )


def normalize_for_compare(text: str) -> str:
    """Lowercase, collapse whitespace, strip markdown formatting characters."""
    text = re.sub(r"[#*_`>\-]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip().lower()


def word_set(text: str) -> set[str]:
    """Bag of words for content-loss measurement."""
    text = re.sub(r"[^\w\s]", " ", text.lower())
    words = text.split()
    return {w for w in words if len(w) >= 3}  # ignore tiny words


def detect_warnings(md_text: str, extracted_text: str, kind: str) -> list[str]:
    """Heuristically detect format issues in the extraction."""
    warnings: list[str] = []

    md_words = word_set(md_text)
    ext_words = word_set(extracted_text)
    if not md_words:
        return warnings

    # --- Content loss ---
    if ext_words:
        retention = len(md_words & ext_words) / len(md_words)
        if retention < 0.7:
            warnings.append(
                f"Major content loss — only {retention:.0%} of unique words from "
                f"resume.md were found in the extracted text. Likely cause: tables "
                f"or complex layouts that PyPDF2 couldn't parse."
            )
        elif retention < 0.85:
            warnings.append(
                f"Some content loss — {retention:.0%} of unique words preserved "
                f"(typical clean PDFs hit 90%+). Worth checking what's missing."
            )

    # --- Section headings ---
    md_headings = re.findall(r"^#+\s+(.+)$", md_text, re.MULTILINE)
    extracted_lower = extracted_text.lower()
    missing_headings = [
        h for h in md_headings
        if normalize_for_compare(h)[:30] not in normalize_for_compare(extracted_lower)
    ]
    if missing_headings:
        warnings.append(
            f"{len(missing_headings)} section heading(s) missing or garbled in "
            f"extracted text: {', '.join(repr(h) for h in missing_headings[:3])}"
            + (f" + {len(missing_headings)-3} more" if len(missing_headings) > 3 else "")
        )

    # --- Repeated content (header/footer leaking into body) ---
    line_counts: dict[str, int] = {}
    for line in extracted_text.splitlines():
        line = line.strip()
        if len(line) < 10:
            continue
        line_counts[line] = line_counts.get(line, 0) + 1
    repeated = [(l, c) for l, c in line_counts.items() if c >= 3]
    if repeated:
        warnings.append(
            f"{len(repeated)} line(s) appear 3+ times — likely header/footer being "
            f"parsed into body content. Example: {repeated[0][0][:60]!r}"
        )

    # --- Unicode / encoding garbage ---
    # Smart quotes, em-dashes, ligatures often come through as garbled bytes.
    weird_chars = re.findall(r"[�   ]", extracted_text)
    if weird_chars:
        warnings.append(
            f"{len(weird_chars)} replacement/control character(s) found — encoding "
            f"issues. Try saving the PDF with embedded fonts or as 'PDF/A'."
        )

    # --- Multi-column heuristic for PDFs ---
    if kind == "pdf":
        # Lines that have a lot of internal whitespace gaps are a sign of
        # column boundaries being preserved in the extracted text.
        big_gaps = sum(
            1 for line in extracted_text.splitlines()
            if re.search(r"\S\s{6,}\S", line)
        )
        if big_gaps > 10:
            warnings.append(
                f"{big_gaps} lines with large internal whitespace gaps — possible "
                f"multi-column layout. ATS reads top-to-bottom, so columns may "
                f"interleave incorrectly."
            )

    # --- DOCX tables ---
    if kind == "docx":
        try:
            from docx import Document
            doc = Document(str(RESUME_DOCX))
            if doc.tables:
                warnings.append(
                    f"{len(doc.tables)} table(s) detected. Many ATSes parse tables "
                    f"poorly — they may flatten columns or jumble row contents. "
                    f"Consider replacing with bullet lists."
                )
        except Exception:
            pass

    return warnings


def main() -> None:
    if not RESUME_MD.exists():
        sys.exit(f"Error: {RESUME_MD} not found. The format check compares your "
                 f"PDF/DOCX against this file, so you need both.")

    resume_file, kind = find_resume_file()

    print(_bar(f"Resume PDF/DOCX format check"))
    print(f"Source:    {RESUME_MD.name}")
    print(f"Extract:   {resume_file.name}  ({kind.upper()})")

    md_text = RESUME_MD.read_text()
    if kind == "pdf":
        extracted = extract_pdf_text(resume_file)
    else:
        extracted = extract_docx_text(resume_file)

    md_chars = len(md_text)
    md_words = len(md_text.split())
    ext_chars = len(extracted)
    ext_words = len(extracted.split())

    print()
    print(f"  resume.md       {md_chars:>7,} chars   {md_words:>5,} words")
    print(f"  {resume_file.name:<15} {ext_chars:>7,} chars   {ext_words:>5,} words")

    if md_words:
        word_pct = ext_words / md_words * 100
        print(f"  word retention: {word_pct:.0f}%")

    warnings = detect_warnings(md_text, extracted, kind)

    print(_bar("Findings"))
    if not warnings:
        print("✓ No issues detected. Your resume should parse cleanly through most ATSes.")
    else:
        for w in warnings:
            print(f"⚠  {w}\n")

    # --- Save the extracted text for manual review ---
    extracted_path = ROOT / f"resume.extracted.{kind}.txt"
    extracted_path.write_text(extracted)
    print(_bar("Extracted text saved"))
    print(f"  {extracted_path.name}")
    print(f"  Open this to see exactly what an ATS would receive.")
    print()


if __name__ == "__main__":
    main()
